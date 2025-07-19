#!/usr/bin/env python3
"""
Word Document Generator Service.

This service handles the generation of Word documents with embedded charts,
tables, and advanced formatting features.
"""

import asyncio
import base64
import io
import json
import logging
import os
import tempfile
import time
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import aiofiles
import aiohttp
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.table import _Cell
from PIL import Image
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend

from app.core.config import settings
from app.models.word_models import (
    Chart,
    ChartType,
    ColorScheme,
    DocumentConfig,
    DocumentSection,
    FontFamily,
    FontStyle,
    JobStatus,
    OutputFormat,
    PageOrientation,
    Table,
    Template
)

logger = logging.getLogger(__name__)


class WordDocumentGenerator:
    """Word document generator with advanced features."""
    
    def __init__(self):
        """Initialize the Word document generator."""
        self.color_schemes = {
            ColorScheme.BLUE: ["#1f77b4", "#aec7e8", "#ff7f0e", "#ffbb78", "#2ca02c"],
            ColorScheme.RED: ["#d62728", "#ff9896", "#ff7f0e", "#ffbb78", "#2ca02c"],
            ColorScheme.GREEN: ["#2ca02c", "#98df8a", "#d62728", "#ff9896", "#ff7f0e"],
            ColorScheme.ORANGE: ["#ff7f0e", "#ffbb78", "#2ca02c", "#98df8a", "#d62728"],
            ColorScheme.PURPLE: ["#9467bd", "#c5b0d5", "#8c564b", "#c49c94", "#e377c2"],
            ColorScheme.TEAL: ["#17becf", "#9edae5", "#bcbd22", "#dbdb8d", "#7f7f7f"],
            ColorScheme.MONOCHROME: ["#000000", "#404040", "#808080", "#c0c0c0", "#ffffff"]
        }
        
        self.font_mapping = {
            FontFamily.CALIBRI: "Calibri",
            FontFamily.ARIAL: "Arial",
            FontFamily.TIMES_NEW_ROMAN: "Times New Roman",
            FontFamily.GEORGIA: "Georgia",
            FontFamily.VERDANA: "Verdana",
            FontFamily.TAHOMA: "Tahoma"
        }
    
    async def generate_document(
        self,
        job_id: int,
        user_id: int,
        document_config: DocumentConfig,
        data_source: Dict[str, Any],
        output_format: OutputFormat = OutputFormat.DOCX
    ) -> Tuple[bool, Optional[str], Optional[str]]:
        """
        Generate a Word document from configuration and data.
        
        Args:
            job_id: Job identifier
            user_id: User identifier
            document_config: Document configuration
            data_source: Data source configuration
            output_format: Output format
            
        Returns:
            Tuple of (success, file_path, error_message)
        """
        start_time = time.time()
        
        try:
            # Update job status to processing
            await self._update_job_status(job_id, JobStatus.PROCESSING, started_at=datetime.utcnow())
            
            logger.info(f"Starting Word document generation for job {job_id}")\n            
            # Fetch data from data source
            data = await self._fetch_data(data_source)
            
            # Create Word document
            document = await self._create_word_document(document_config, data, job_id)
            
            # Generate file path
            file_name = f"document_{job_id}_{int(time.time())}.{output_format.value}"
            file_path = os.path.join(settings.WORD_OUTPUT_DIR, file_name)
            
            # Save document
            await self._save_document(document, file_path)
            
            # Calculate metadata
            file_size = os.path.getsize(file_path)
            page_count = await self._count_pages(file_path)
            chart_count = len(document_config.charts)
            table_count = len(document_config.tables)
            section_count = len(document_config.layout.sections)
            generation_time_ms = int((time.time() - start_time) * 1000)
            
            # Update job with success
            await self._update_job_completion(
                job_id=job_id,
                status=JobStatus.COMPLETED,
                file_path=file_path,
                file_size=file_size,
                page_count=page_count,
                chart_count=chart_count,
                table_count=table_count,
                section_count=section_count,
                generation_time_ms=generation_time_ms
            )
            
            logger.info(f"Word document generation completed for job {job_id}")
            return True, file_path, None
            
        except Exception as e:
            error_message = str(e)
            logger.error(f"Word document generation failed for job {job_id}: {error_message}")
            
            # Update job with failure
            await self._update_job_status(
                job_id,
                JobStatus.FAILED,
                error_message=error_message,
                completed_at=datetime.utcnow()
            )
            
            return False, None, error_message
    
    async def _create_word_document(
        self,
        config: DocumentConfig,
        data: Dict[str, Any],
        job_id: int
    ) -> Document:
        """Create Word document from configuration and data."""
        # Create new document
        document = Document()
        
        # Set document properties
        self._set_document_properties(document, config.metadata)
        
        # Set page setup
        self._set_page_setup(document, config.layout.page_setup)
        
        # Set default font
        self._set_default_styles(document, config.default_font, config.template)
        
        # Add header and footer
        if config.layout.header:
            self._add_header(document, config.layout.header)
        
        if config.layout.footer:
            self._add_footer(document, config.layout.footer)
        
        # Add watermark
        if config.layout.watermark:
            self._add_watermark(document, config.layout.watermark)
        
        # Add title page
        await self._add_title_page(document, config)
        
        # Add table of contents if requested
        if config.layout.table_of_contents:
            self._add_table_of_contents(document)
        
        # Process sections
        for section in config.layout.sections:
            await self._add_section(document, section, config, data)
        
        return document
    
    def _set_document_properties(self, document: Document, metadata):
        """Set document properties and metadata."""
        core_props = document.core_properties
        core_props.title = metadata.title
        core_props.subject = metadata.subject
        core_props.author = metadata.author
        core_props.category = metadata.category
        core_props.comments = metadata.comments
        core_props.keywords = ", ".join(metadata.keywords) if metadata.keywords else ""
        core_props.created = datetime.utcnow()
        core_props.modified = datetime.utcnow()
    
    def _set_page_setup(self, document: Document, page_setup):
        """Set page setup including orientation and margins."""
        section = document.sections[0]
        
        # Set page orientation
        if page_setup.orientation == PageOrientation.LANDSCAPE:
            section.orientation = WD_ORIENT.LANDSCAPE
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
        
        # Set margins
        section.top_margin = Inches(page_setup.margins.top)
        section.bottom_margin = Inches(page_setup.margins.bottom)
        section.left_margin = Inches(page_setup.margins.left)
        section.right_margin = Inches(page_setup.margins.right)
    
    def _set_default_styles(self, document: Document, font_style: FontStyle, template: Template):
        """Set default document styles based on template and font settings."""
        styles = document.styles
        
        # Modify Normal style
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = self.font_mapping.get(font_style.family, "Calibri")
        normal_font.size = Pt(font_style.size)
        normal_font.bold = font_style.bold
        normal_font.italic = font_style.italic
        
        # Set font color
        if font_style.color != "#000000":
            rgb_color = self._hex_to_rgb(font_style.color)
            normal_font.color.rgb = RGBColor(rgb_color[0], rgb_color[1], rgb_color[2])
        
        # Create or modify heading styles based on template
        self._create_heading_styles(document, template, font_style)
    
    def _create_heading_styles(self, document: Document, template: Template, base_font: FontStyle):
        """Create heading styles based on template."""
        styles = document.styles
        
        # Template-specific styling
        if template == Template.PROFESSIONAL:
            colors = ["#2E74B5", "#1F4E79", "#0F2C4C"]
        elif template == Template.CORPORATE:
            colors = ["#C5504B", "#943634", "#6B1F1C"]
        elif template == Template.ACADEMIC:
            colors = ["#70AD47", "#548235", "#375623"]
        elif template == Template.REPORT:
            colors = ["#FFC000", "#D99B00", "#B8860B"]
        else:  # MINIMAL
            colors = ["#404040", "#202020", "#000000"]
        
        # Update heading styles
        for i, (style_name, color) in enumerate(zip(['Heading 1', 'Heading 2', 'Heading 3'], colors)):
            if style_name in styles:
                heading_style = styles[style_name]
            else:
                heading_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            
            heading_font = heading_style.font
            heading_font.name = self.font_mapping.get(base_font.family, "Calibri")
            heading_font.size = Pt(base_font.size + 4 - i)
            heading_font.bold = True
            
            rgb_color = self._hex_to_rgb(color)
            heading_font.color.rgb = RGBColor(rgb_color[0], rgb_color[1], rgb_color[2])
    
    def _add_header(self, document: Document, header_config):
        """Add header to document."""
        section = document.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        
        # Set text and alignment
        header_para.text = header_config.text
        
        if header_config.alignment == "center":
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif header_config.alignment == "right":
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Apply font styling
        self._apply_font_style(header_para.runs[0] if header_para.runs else header_para.add_run(), header_config.font_style)
        
        # Add page number and date if requested
        if header_config.include_page_number or header_config.include_date:
            additional_text = []
            if header_config.include_date:
                additional_text.append(datetime.now().strftime("%B %d, %Y"))
            if header_config.include_page_number:
                additional_text.append("Page ")
            
            if additional_text:
                header_para.add_run(" | " + " | ".join(additional_text))
    
    def _add_footer(self, document: Document, footer_config):
        """Add footer to document."""
        section = document.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        
        # Set text and alignment
        footer_para.text = footer_config.text
        
        if footer_config.alignment == "center":
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif footer_config.alignment == "right":
            footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Apply font styling
        self._apply_font_style(footer_para.runs[0] if footer_para.runs else footer_para.add_run(), footer_config.font_style)
        
        # Add page number and date if requested
        if footer_config.include_page_number or footer_config.include_date:
            additional_text = []
            if footer_config.include_date:
                additional_text.append(datetime.now().strftime("%B %d, %Y"))
            if footer_config.include_page_number:
                additional_text.append("Page ")
            
            if additional_text:
                footer_para.add_run(" | " + " | ".join(additional_text))
    
    def _add_watermark(self, document: Document, watermark_text: str):
        """Add watermark to document."""
        # This is a simplified watermark implementation
        # In a production environment, you might want to use more sophisticated watermarking
        for section in document.sections:
            header = section.header
            watermark_para = header.add_paragraph()
            watermark_run = watermark_para.add_run(watermark_text)
            watermark_run.font.color.rgb = RGBColor(200, 200, 200)  # Light gray
            watermark_run.font.size = Pt(48)
            watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    async def _add_title_page(self, document: Document, config: DocumentConfig):
        """Add title page to document."""
        # Add title
        title_para = document.add_paragraph()
        title_run = title_para.add_run(config.metadata.title)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some spacing
        document.add_paragraph()
        
        # Add subject if available
        if config.metadata.subject:
            subject_para = document.add_paragraph()
            subject_run = subject_para.add_run(config.metadata.subject)
            subject_run.font.size = Pt(16)
            subject_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add author and company
        if config.metadata.author or config.metadata.company:
            document.add_paragraph()
            author_para = document.add_paragraph()
            if config.metadata.author:
                author_para.add_run(f"Author: {config.metadata.author}")
            if config.metadata.company:
                if config.metadata.author:
                    author_para.add_run("\\n")
                author_para.add_run(f"Company: {config.metadata.company}")
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_para = document.add_paragraph()
        date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page break
        document.add_page_break()
    
    def _add_table_of_contents(self, document: Document):
        """Add table of contents to document."""
        # Add TOC title
        toc_para = document.add_paragraph()
        toc_run = toc_para.add_run("Table of Contents")
        toc_run.font.size = Pt(16)
        toc_run.font.bold = True
        toc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add placeholder for TOC
        # In a real implementation, you would populate this with actual headings
        toc_content = document.add_paragraph()
        toc_content.add_run("Table of Contents will be generated here...")
        
        # Add page break
        document.add_page_break()
    
    async def _add_section(
        self,
        document: Document,
        section: DocumentSection,
        config: DocumentConfig,
        data: Dict[str, Any]
    ):
        """Add a section to the document."""
        # Add section title if provided
        if section.title:
            heading_para = document.add_heading(section.title, level=1)
        
        # Process content based on type
        if section.content_type == "text" and section.text_content:
            para = document.add_paragraph(section.text_content)
            
        elif section.content_type == "chart" and section.content_id:
            # Find and add chart
            chart = next((c for c in config.charts if c.id == section.content_id), None)
            if chart:
                await self._add_chart(document, chart, config.color_scheme)
        
        elif section.content_type == "table" and section.content_id:
            # Find and add table
            table = next((t for t in config.tables if t.id == section.content_id), None)
            if table:
                self._add_table(document, table)
        
        elif section.content_type == "page_break":
            document.add_page_break()
        
        # Add some spacing after section
        document.add_paragraph()
    
    async def _add_chart(self, document: Document, chart: Chart, color_scheme: ColorScheme):
        """Add chart to document as image."""
        try:
            # Generate chart image using matplotlib
            chart_image = await self._generate_chart_image(chart, color_scheme)
            
            # Add chart image to document
            if chart_image:
                para = document.add_paragraph()
                run = para.add_run()
                run.add_picture(chart_image, width=Inches(chart.config.width / 100), height=Inches(chart.config.height / 100))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add chart title if provided
                if chart.config.title:
                    title_para = document.add_paragraph()
                    title_run = title_para.add_run(chart.config.title)
                    title_run.font.bold = True
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        except Exception as e:
            logger.error(f"Failed to add chart {chart.id}: {e}")
            # Add placeholder text
            para = document.add_paragraph(f"[Chart: {chart.config.title or chart.id}]")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    async def _generate_chart_image(self, chart: Chart, color_scheme: ColorScheme) -> Optional[io.BytesIO]:
        """Generate chart image using matplotlib."""
        try:
            # Set up matplotlib
            plt.style.use('default')
            fig, ax = plt.subplots(figsize=(chart.config.width / 100, chart.config.height / 100))
            
            # Get colors
            colors = self.color_schemes.get(color_scheme, self.color_schemes[ColorScheme.BLUE])
            
            # Extract data
            labels = chart.data.labels
            datasets = chart.data.datasets
            
            # Generate chart based on type
            if chart.config.chart_type == ChartType.BAR:
                self._create_bar_chart(ax, labels, datasets, colors, chart.config)
            elif chart.config.chart_type == ChartType.LINE:
                self._create_line_chart(ax, labels, datasets, colors, chart.config)
            elif chart.config.chart_type == ChartType.PIE:
                self._create_pie_chart(ax, labels, datasets[0], colors, chart.config)
            elif chart.config.chart_type == ChartType.AREA:
                self._create_area_chart(ax, labels, datasets, colors, chart.config)
            elif chart.config.chart_type == ChartType.SCATTER:
                self._create_scatter_chart(ax, labels, datasets, colors, chart.config)
            else:
                # Default to bar chart
                self._create_bar_chart(ax, labels, datasets, colors, chart.config)
            
            # Set title and labels
            if chart.config.title:
                ax.set_title(chart.config.title, fontsize=14, fontweight='bold')
            
            # Show/hide grid
            ax.grid(chart.config.show_grid, alpha=0.3)
            
            # Show/hide legend
            if chart.config.show_legend and len(datasets) > 1:
                ax.legend()
            
            # Save to BytesIO
            img_buffer = io.BytesIO()
            plt.savefig(img_buffer, format='PNG', dpi=300, bbox_inches='tight')
            plt.close(fig)
            
            img_buffer.seek(0)
            return img_buffer
            
        except Exception as e:
            logger.error(f"Failed to generate chart image: {e}")
            return None
    
    def _create_bar_chart(self, ax, labels, datasets, colors, config):
        """Create bar chart."""
        x = range(len(labels))
        width = 0.8 / len(datasets) if len(datasets) > 1 else 0.8
        
        for i, dataset in enumerate(datasets):
            offset = (i - len(datasets) / 2 + 0.5) * width
            ax.bar([xi + offset for xi in x], dataset["data"], width, 
                   label=dataset["label"], color=colors[i % len(colors)])
        
        ax.set_xticks(x)
        ax.set_xticklabels(labels)
    
    def _create_line_chart(self, ax, labels, datasets, colors, config):
        """Create line chart."""
        for i, dataset in enumerate(datasets):
            ax.plot(labels, dataset["data"], label=dataset["label"], 
                   color=colors[i % len(colors)], marker='o')
    
    def _create_pie_chart(self, ax, labels, dataset, colors, config):
        """Create pie chart."""
        ax.pie(dataset["data"], labels=labels, colors=colors[:len(labels)], 
               autopct='%1.1f%%', startangle=90)
        ax.axis('equal')
    
    def _create_area_chart(self, ax, labels, datasets, colors, config):
        """Create area chart."""
        for i, dataset in enumerate(datasets):
            ax.fill_between(labels, dataset["data"], alpha=0.7, 
                           label=dataset["label"], color=colors[i % len(colors)])
    
    def _create_scatter_chart(self, ax, labels, datasets, colors, config):
        """Create scatter chart."""
        for i, dataset in enumerate(datasets):
            # For scatter plot, assume first dataset is x-values, second is y-values
            if i == 0:
                x_data = dataset["data"]
            else:
                ax.scatter(x_data, dataset["data"], label=dataset["label"], 
                          color=colors[i % len(colors)])
    
    def _add_table(self, document: Document, table_data: Table):
        """Add table to document."""
        try:
            # Create table
            table = document.add_table(rows=1, cols=len(table_data.config.columns))
            table.style = 'Table Grid'
            
            # Add header row
            header_row = table.rows[0]
            for i, column in enumerate(table_data.config.columns):
                cell = header_row.cells[i]
                cell.text = column.label
                
                # Apply header styling
                self._apply_cell_styling(cell, table_data.config.header_style)
            
            # Add data rows
            for row_data in table_data.data:
                row = table.add_row()
                for i, column in enumerate(table_data.config.columns):
                    cell = row.cells[i]
                    value = row_data.get(column.name, "")
                    cell.text = str(value)
                    
                    # Apply cell alignment
                    if column.alignment == "center":
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif column.alignment == "right":
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    # Apply row styling
                    self._apply_cell_styling(cell, table_data.config.row_style)
        
        except Exception as e:
            logger.error(f"Failed to add table {table_data.id}: {e}")
            # Add placeholder text
            para = document.add_paragraph(f"[Table: {table_data.id}]")
    
    def _apply_cell_styling(self, cell: _Cell, style_config: Dict[str, Any]):
        """Apply styling to table cell."""
        if not style_config:
            return
        
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if "font_size" in style_config:
                    run.font.size = Pt(style_config["font_size"])
                if "bold" in style_config:
                    run.font.bold = style_config["bold"]
                if "italic" in style_config:
                    run.font.italic = style_config["italic"]
                if "color" in style_config:
                    rgb_color = self._hex_to_rgb(style_config["color"])
                    run.font.color.rgb = RGBColor(rgb_color[0], rgb_color[1], rgb_color[2])
    
    def _apply_font_style(self, run, font_style: FontStyle):
        """Apply font style to a run."""
        run.font.name = self.font_mapping.get(font_style.family, "Calibri")
        run.font.size = Pt(font_style.size)
        run.font.bold = font_style.bold
        run.font.italic = font_style.italic
        
        if font_style.color != "#000000":
            rgb_color = self._hex_to_rgb(font_style.color)
            run.font.color.rgb = RGBColor(rgb_color[0], rgb_color[1], rgb_color[2])
    
    def _hex_to_rgb(self, hex_color: str) -> tuple:
        """Convert hex color to RGB tuple."""
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    
    async def _fetch_data(self, data_source: Dict[str, Any]) -> Dict[str, Any]:
        """Fetch data from data source."""
        if "static_source" in data_source:
            return data_source["static_source"].get("data", {})
        elif "query_source" in data_source:
            # In a real implementation, this would execute the query
            return {"query_result": "mock_data"}
        elif "file_source" in data_source:
            # In a real implementation, this would load from file
            return {"file_data": "mock_data"}
        else:
            return {}
    
    async def _save_document(self, document: Document, file_path: str) -> None:
        """Save Word document to file."""
        document.save(file_path)
    
    async def _count_pages(self, file_path: str) -> int:
        """Count pages in Word document."""
        # This is a simplified implementation
        # In a real implementation, you might use python-docx2txt or similar
        try:
            doc = Document(file_path)
            # Rough estimation based on content
            total_elements = len(doc.paragraphs) + sum(len(table.rows) for table in doc.tables)
            # Assume ~50 elements per page (very rough estimate)
            return max(1, total_elements // 50)
        except Exception:
            return 1
    
    async def _update_job_status(
        self,
        job_id: int,
        status: JobStatus,
        error_message: Optional[str] = None,
        started_at: Optional[datetime] = None,
        completed_at: Optional[datetime] = None
    ) -> None:
        """Update job status in database."""
        try:
            from app.core.database import get_db_session, update_job_status
            
            async with get_db_session() as session:
                await update_job_status(
                    session,
                    job_id,
                    status.value,
                    error_message,
                    started_at,
                    completed_at
                )
            logger.info(f"Job {job_id} status updated to {status.value}")
        except Exception as e:
            logger.error(f"Failed to update job status: {e}")
    
    async def _update_job_completion(
        self,
        job_id: int,
        status: JobStatus,
        file_path: str,
        file_size: int,
        page_count: int,
        chart_count: int,
        table_count: int,
        section_count: int,
        generation_time_ms: int
    ) -> None:
        """Update job with completion details."""
        try:
            from app.core.database import get_db_session, update_job_completion
            
            async with get_db_session() as session:
                await update_job_completion(
                    session,
                    job_id,
                    status.value,
                    file_path,
                    file_size,
                    page_count,
                    chart_count,
                    table_count,
                    section_count,
                    generation_time_ms
                )
            logger.info(f"Job {job_id} completed with {chart_count} charts, {table_count} tables, {section_count} sections")
        except Exception as e:
            logger.error(f"Failed to update job completion: {e}")


# Global generator instance
word_document_generator = WordDocumentGenerator()


# Export commonly used functions
__all__ = ["word_document_generator", "WordDocumentGenerator"]